from data import ORIGENS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import ast
import os


# =============================================================================
# ELEMENTOS DO RELATÓRIO
# =============================================================================

def criar_header(report, logo):
    """Cria o header com o logotipo e o título do relatório."""
    header = report.sections[0].header
    header_paragraph = header.paragraphs[0]
    run_logo = header_paragraph.add_run()
    try:
        if logo and os.path.exists(logo):
            run_logo.add_picture(logo, width=Inches(1.2))
    except Exception as e:
        print(f"[!] Erro ao inserir logo: {e}")

    run_text = header_paragraph.add_run(
        "\tRelatório de Justificativas de Falsos Positivos\tCONFIDENCIAL"
    )
    run_text.font.size = Pt(10)
    run_text.font.bold = True
    run_text.font.color.rgb = RGBColor(0, 0, 0)


def criar_footer(report):
    """Cria o rodapé (legal notice)."""
    footer = report.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    run = footer_paragraph.add_run(
        "\tEsse documento foi classificado pelo time de Segurança. "
        "O acesso está autorizado exclusivamente aos colaboradores envolvidos no processo.\t"
    )
    run.font.size = Pt(6)
    run.font.italic = True


def perguntas_dict(perguntas, respostas):
    """Cria um dicionário associando perguntas e respostas."""
    return dict(zip(perguntas, respostas))


def id_projeto(report, origem, respostas):
    """Cria a seção de identificação do projeto com tabela baseada na origem."""
    report.add_heading("Descrição do Projeto", level=1)
    report.add_paragraph()

    # 🔹 Usa o bloco de dados da aplicação correspondente à origem
    origem_cfg = ORIGENS.get(origem, {})
    campos_dados = origem_cfg.get("dados_aplicacao", [])

    # Extrai apenas os labels definidos no bloco de dados-aplicação
    perguntas = [campo["label"] for campo in campos_dados]

    # Cria dicionário associando perguntas ↔ respostas
    perguntas_respostas = perguntas_dict(perguntas, respostas)

    # Cria tabela como antes
    table = report.add_table(rows=len(perguntas_respostas), cols=2, style="Light Grid Accent 2")

    for i, (pergunta, resposta) in enumerate(perguntas_respostas.items()):
        table.cell(i, 0).text = pergunta
        table.cell(i, 1).text = resposta if resposta else "—"


def sessao_justificativa(report):
    """Inicia a seção de justificativas."""
    report.add_paragraph()
    report.add_heading("Vulnerabilidades Justificadas", level=1)


def tabela_justificativa(report, justificativa):
    """Cria uma tabela para uma justificativa (dict)."""
    report.add_paragraph()

    table = report.add_table(rows=len(justificativa) + 2, cols=2, style="Light Grid Accent 2")
    perguntas_col = table.columns[0].cells
    respostas_col = table.columns[1].cells

    row = 0
    for chave, valor in justificativa.items():
        perguntas_col[row].text = chave

        # ==========================================================
        # TRATAMENTO ROBUSTO DE EVIDÊNCIAS
        # ==========================================================
        if chave == "Evidência":
            arquivos = []
            if isinstance(valor, list):
                arquivos = [p for p in valor if isinstance(p, str) and p.strip()]
            elif isinstance(valor, str) and valor.strip():
                arquivos = [valor]

            if not arquivos:
                respostas_col[row].text = "Sem evidências anexadas"
            else:
                par = respostas_col[row].paragraphs[0]
                for image_path in arquivos:
                    try:
                        if os.path.exists(image_path):
                            par.add_run().add_picture(image_path, width=Inches(4.5))
                        else:
                            par.add_run(f"[Arquivo não encontrado: {image_path}]")
                    except Exception as e:
                        par.add_run(f"[Erro ao inserir imagem: {e}]")
        else:
            respostas_col[row].text = (
                valor[0] if isinstance(valor, list) and len(valor) == 1 else str(valor)
            )

        row += 1

    perguntas_col[row].text = "Comentário DSS:"
    perguntas_col[row + 1].text = "Situação [Aprovado/Reprovado]:"


def salvar_relatorio(report, nome_aplicacao):
    """Salva o relatório Word."""
    safe_name = nome_aplicacao.replace(" ", "_") or "relatorio"
    output_path = f"{safe_name}_Relatorio_Justificativa.docx"
    report.save(output_path)
    print(f"[+] Relatório salvo em: {output_path}")


# =============================================================================
# CONVERSÃO DE DADOS
# =============================================================================

def tuple_to_list(tuple_justificativa):
    """Converte o tuple do listbox (strings de dicts) em lista de dicionários reais."""
    justificativa_lista = []
    for item in tuple_justificativa:
        try:
            data_dict = ast.literal_eval(item)
            justificativa_lista.append(data_dict)
        except Exception as e:
            print(f"[!] Erro ao converter item: {e}")
    return justificativa_lista


# =============================================================================
# PIPELINE PRINCIPAL
# =============================================================================

def criar_report(logo, respostas, tuple_justificativa):
    """Função principal para gerar o relatório completo."""
    # 🔹 Extraímos a origem do primeiro item (todas justificativas têm a mesma origem)
    justificativas = tuple_to_list(tuple_justificativa)
    origem = justificativas[0].get("Origem", "Desconhecida") if justificativas else "Desconhecida"

    report = Document()
    criar_header(report, logo)
    id_projeto(report, origem, respostas)
    sessao_justificativa(report)

    for j in justificativas:
        tabela_justificativa(report, dict(j))

    criar_footer(report)
    salvar_relatorio(report, respostas[0])
